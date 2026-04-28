import os
import re
import json
import random
from docx import Document
from openai import OpenAI

# ================== 配置区域 ==================
POLICY_FOLDER = "./policies"
OUTPUT_POLICY_JSON = "policy_library.json"
OUTPUT_PACR_DATA = "pacr_training_data.json"

DEEPSEEK_API_KEY = "sk-daef1804b9cd447f8b925ca2e6e1c9a3"
DEEPSEEK_BASE_URL = "https://api.deepseek.com"

QA_PER_ACTIVE_POLICY = 15
QA_PER_REPEALED_POLICY = 5
NEGATIVE_SAMPLE_COUNT = 4

# ================== 工具函数 ==================
def clean_text(text):
    return re.sub(r'\s+', ' ', text).strip()

def extract_doc_no(text):
    m = re.search(r'国家税务总局令第\s*(\d+)\s*号', text)
    if m:
        return f"国家税务总局令第{m.group(1)}号"
    return ""

def extract_effective_date(text):
    for pat in [
        r'自\s*(\d{4})年(\d{1,2})月(\d{1,2})日\s*起施行',
        r'自\s*(\d{4})年(\d{1,2})月(\d{1,2})日\s*施行',
        r'(\d{4})年(\d{1,2})月(\d{1,2})日\s*起施行'
    ]:
        m = re.search(pat, text)
        if m:
            y, mth, d = m.groups()
            return f"{y}-{mth.zfill(2)}-{d.zfill(2)}"
    return ""

def extract_title_from_doc(doc):
    """从文档的前5个段落中提取标题，避免混入正文"""
    candidates = []
    for para in doc.paragraphs[:5]:
        text = clean_text(para.text)
        if not text:
            continue
        # 标题通常含有这些关键词
        if any(kw in text for kw in ['管理办法', '管理办法（', '规定', '决定', '条例', '办法']):
            # 排除文号行
            if '国家税务总局令' in text and '号' in text and len(text) < 30:
                continue
            candidates.append((len(text), text))
    if candidates:
        # 选择最长的候选作为标题
        return sorted(candidates, key=lambda x: x[0], reverse=True)[0][1]
    return ""

# ================== 第一步：政策解析 ==================
def parse_policy_docx(file_path):
    doc = Document(file_path)
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    title = extract_title_from_doc(doc)
    if not title:
        # 从文件名提取
        title = os.path.basename(file_path).replace('.docx','').replace('.doc','')
        print(f"  ⚠ 未能自动提取标题，使用文件名: {title}")

    doc_no = extract_doc_no(full_text)
    effective_date = extract_effective_date(full_text)

    # 提取条款
    articles = []
    current = None
    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue
        if re.match(r'^第\s*[一二三四五六七八九十百千万\d]+\s*条', text):
            if current:
                articles.append(current)
            current = {"heading": text, "content": ""}
        elif current:
            current["content"] += text + "\n"
    if current:
        articles.append(current)

    return {
        "source_file": os.path.basename(file_path),
        "title": title,
        "doc_no": doc_no,
        "effective_date": effective_date,
        "full_text": full_text[:2000] + "..." if len(full_text) > 2000 else full_text,
        "articles": articles,
        "article_count": len(articles),
        "status": "active",
        "repealed_by": None,
        "repeal_date": None
    }

def parse_all_policies(folder):
    policies = []
    for fname in os.listdir(folder):
        if fname.endswith('.docx'):
            path = os.path.join(folder, fname)
            try:
                p = parse_policy_docx(path)
                policies.append(p)
                print(f"✓ 解析: {fname} -> {p['article_count']}条, 标题: {p['title']}")
            except Exception as e:
                print(f"✗ 解析失败 {fname}: {e}")
    return policies

# ================== 第二步：废止关系识别 ==================
def detect_repeal_relationships(policies):
    docno_to_idx = {}
    for i, p in enumerate(policies):
        clean_no = re.sub(r'\s+', '', p['doc_no'])
        if clean_no:
            docno_to_idx[clean_no] = i

    for p in policies:
        if '废止' in p['title']:
            pattern = r'《([^》]+)》[（(]国家税务总局令第(\d+)号[）)]'
            matches = re.findall(pattern, p['full_text'])
            for m in matches:
                target_doc_no = f"国家税务总局令第{m[1]}号"
                clean_target = re.sub(r'\s+', '', target_doc_no)
                if clean_target in docno_to_idx:
                    idx = docno_to_idx[clean_target]
                    policies[idx]['status'] = 'repealed'
                    policies[idx]['repealed_by'] = p['doc_no']
                    policies[idx]['repeal_date'] = p['effective_date']
                    print(f"✓ 标记废止: {policies[idx]['title']} 被 {p['title']} 废止")

# ================== 第三步：问答合成 ==================
client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url=DEEPSEEK_BASE_URL)

def generate_qa_for_policy(policy):
    if policy['status'] == 'repealed':
        num = QA_PER_REPEALED_POLICY
        prompt = f"""你是一位税务政策历史研究者。以下政策已经废止，请生成{num}个关于该政策历史沿革的问题。
要求：问题包含明确时间指向；答案首句注明废止信息。
政策标题：{policy['title']}
文号：{policy['doc_no']}
废止日期：{policy.get('repeal_date', '未知')}
内容摘要：{policy['full_text'][:800]}

请以JSON数组格式输出：[{{"question":"...", "answer":"...", "difficulty":"easy/medium/hard", "type":"历史追溯/...", "reference_article":"第X条"}}]
"""
    else:
        num = QA_PER_ACTIVE_POLICY
        prompt = f"""你是一位税务政策问答专家。请根据以下现行有效政策生成{num}个高质量问答对。
要求：问题类型多样化。
政策标题：{policy['title']}
文号：{policy['doc_no']}
生效日期：{policy['effective_date']}
内容摘要：{policy['full_text'][:800]}

请以JSON数组格式输出：[{{"question":"...", "answer":"...", "difficulty":"easy/medium/hard", "type":"适用条件/办理流程/...", "reference_article":"第X条"}}]
"""
    try:
        resp = client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=3000
        )
        content = resp.choices[0].message.content
        json_match = re.search(r'\[.*\]', content, re.DOTALL)
        if json_match:
            qa_list = json.loads(json_match.group())
        else:
            qa_list = []
        for qa in qa_list:
            qa['policy_id'] = policy['doc_no']
            qa['policy_title'] = policy['title']
            qa['policy_status'] = policy['status']
        return qa_list
    except Exception as e:
        print(f"  生成失败: {e}")
        return []

def generate_all_qa(policies):
    all_qa = []
    for p in policies:
        print(f"为《{p['title']}》生成问答...")
        qa = generate_qa_for_policy(p)
        all_qa.extend(qa)
        print(f"  → 生成 {len(qa)} 条")
    return all_qa

# ================== 第四步：构建 PACR 训练样本 ==================
def build_pacr_samples(qa_pairs, policies):
    samples = []
    for qa in qa_pairs:
        pos_id = qa['policy_id']
        neg_pool = [p for p in policies if p['doc_no'] != pos_id]
        if len(neg_pool) == 0:
            # 只有一条政策，无法构建对比样本，跳过
            continue
        if len(neg_pool) < NEGATIVE_SAMPLE_COUNT:
            neg_samples = random.choices(neg_pool, k=NEGATIVE_SAMPLE_COUNT)
        else:
            neg_samples = random.sample(neg_pool, NEGATIVE_SAMPLE_COUNT)

        sample = {
            "question": qa['question'],
            "answer": qa['answer'],
            "positive_policy": {
                "id": pos_id,
                "title": qa['policy_title'],
                "text": next((p['full_text'] for p in policies if p['doc_no'] == pos_id), "")
            },
            "negative_policies": [
                {"id": n['doc_no'], "title": n['title'], "text": n['full_text']}
                for n in neg_samples
            ],
            "difficulty": qa.get('difficulty', 'medium'),
            "type": qa.get('type', '其他')
        }
        samples.append(sample)
    return samples

# ================== 主流程 ==================
def main():
    print("===== 阶段1：解析政策文件 =====")
    policies = parse_all_policies(POLICY_FOLDER)
    if len(policies) < 2:
        print("❌ 政策文件太少，请确保所有.docx文件已放入 policies 文件夹。")
        return
    detect_repeal_relationships(policies)
    with open(OUTPUT_POLICY_JSON, 'w', encoding='utf-8') as f:
        json.dump(policies, f, ensure_ascii=False, indent=2)
    print(f"政策库已保存至 {OUTPUT_POLICY_JSON}，共 {len(policies)} 条\n")

    print("===== 阶段2：生成问答对（调用 DeepSeek）=====")
    qa_pairs = generate_all_qa(policies)
    print(f"共生成 {len(qa_pairs)} 个问答对\n")

    print("===== 阶段3：构建 PACR 训练样本（纯对比学习，无人工权重）=====")
    pacr_data = build_pacr_samples(qa_pairs, policies)
    with open(OUTPUT_PACR_DATA, 'w', encoding='utf-8') as f:
        json.dump(pacr_data, f, ensure_ascii=False, indent=2)
    print(f"PACR 训练数据已保存至 {OUTPUT_PACR_DATA}，共 {len(pacr_data)} 个样本")
    print("\n===== 数据准备完成！全程未使用任何人工权重。 =====")

if __name__ == "__main__":
    main()